#!/usr/bin/env python3
"""
LLM-Powered Intelligent Document Query System - FastAPI Server
Refactored for LlamaIndex with a focus on performance optimization.
"""
import logging
import os
import sys
import requests
import tempfile
import asyncio
import uvicorn
import time
import docx
import json
from typing import Optional, List, Dict, Any
from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Security
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
from contextlib import asynccontextmanager
from urllib.parse import urlparse

# Import libraries for new document types and advanced parsing
import PyPDF2
from email.message import Message
import pypdf
import fitz # PyMuPDF for robust PDF and image handling
from tabulate import tabulate # For converting tables to markdown

# Ensure the project structure is correct for imports
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from config import config
from src.query_processor import QueryProcessor
from llama_index.core import VectorStoreIndex, Settings
from llama_index.core.ingestion import IngestionPipeline
from llama_index.core.storage.storage_context import StorageContext
from llama_index.core.node_parser import SentenceSplitter, SentenceWindowNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.vector_stores.pinecone import PineconeVectorStore
from pinecone import Pinecone
from llama_index.llms.openai import OpenAI
from llama_index.core.schema import Document

# Configure logging
logging.basicConfig(
    level=getattr(logging, config.LOG_LEVEL),
    format='%(asctime)ss - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# --- Pydantic models ---
class DecisionResponse(BaseModel):
    decision: str
    amount: Optional[float] = None
    justification: str
    clauses: List[str]

class QueryRequest(BaseModel):
    query: str
    document_path: Optional[str] = None

class BatchQueryRequest(BaseModel):
    documents: Optional[str] = Field(None, description="URL to document for indexing")
    questions: List[str] = Field(..., description="List of questions to answer")

class HackRxResponse(BaseModel):
    answers: List[str] = Field(..., description="List of answers to the questions")

# Global system components
query_processor: QueryProcessor = QueryProcessor()
pipeline: Optional[IngestionPipeline] = None
index: Optional[VectorStoreIndex] = None
executor = ThreadPoolExecutor(max_workers=os.cpu_count() * 2 or 8) 
pc_client = Pinecone(api_key=config.PINECONE_API_KEY)

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Handles startup and shutdown events for the application."""
    global index, pipeline
    logger.info("🚀 System initialization started on startup.")
    Settings.embed_model = OpenAIEmbedding(model=config.EMBEDDING_MODEL)
    Settings.llm = OpenAI(model=config.OPENAI_MODEL)
    Settings.chunk_size = config.CHUNK_SIZE
    Settings.chunk_overlap = config.CHUNK_OVERLAP

    try:
        existing_indexes = pc_client.list_indexes().names()
        if config.PINECONE_INDEX_NAME not in existing_indexes:
            logger.info(f"🔍 Pinecone index '{config.PINECONE_INDEX_NAME}' not found. Creating a new one...")
            pc_client.create_index(
                name=config.PINECONE_INDEX_NAME,
                dimension=1536,
                metric="cosine",
                spec={"serverless": {"cloud": "aws", "region": config.PINECONE_ENVIRONMENT}}
            )
            logger.info(f"✅ Pinecone index '{config.PINECONE_INDEX_NAME}' created successfully.")
        else:
            logger.info(f"✅ Pinecone index '{config.PINECONE_INDEX_NAME}' already exists. Connecting to it.")

        pinecone_vector_store = PineconeVectorStore(
            pinecone_client=pc_client,
            index_name=config.PINECONE_INDEX_NAME,
            environment=config.PINECONE_ENVIRONMENT
        )
        storage_context = StorageContext.from_defaults(vector_store=pinecone_vector_store)
        index = VectorStoreIndex([], storage_context=storage_context, show_progress=True)
        
        # Corrected node parser configuration, fixing the TypeError
        node_parser = SentenceWindowNodeParser.from_defaults(
            window_size=3,
            window_metadata_key="window",
            original_text_metadata_key="original_text",
            sentence_splitter=SentenceSplitter(
                chunk_size=config.CHUNK_SIZE,
                chunk_overlap=config.CHUNK_OVERLAP
            )
        )

        pipeline = IngestionPipeline(
            transformations=[
                SentenceSplitter(chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP),
                Settings.embed_model
            ],
            vector_store=pinecone_vector_store
        )
        logger.info("✨ LlamaIndex and Pinecone components initialized.")
        yield
    except Exception as e:
        logger.error(f"❌ System initialization failed: {e}")
        raise
    finally:
        logger.info("🛑 Shutting down thread pool executor...")
        executor.shutdown(wait=True)
        logger.info("🛑 Thread pool executor shut down.")

app = FastAPI(
    title="LLM Document Processing System",
    description="Intelligent query-retrieval system for unstructured documents, powered by LlamaIndex.",
    version="2.1.3",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

security = HTTPBearer(auto_error=False)

def verify_token(credentials: Optional[HTTPAuthorizationCredentials] = Security(security)):
    """Verify Bearer token authentication."""
    SECRET_TOKEN = "hackrx-test-token-2026"
    if not credentials or credentials.credentials != SECRET_TOKEN:
        raise HTTPException(status_code=401, detail="Invalid or missing authentication token")
    return credentials.credentials

@app.get("/")
async def root():
    return {
        "message": "LLM Document Processing System",
        "version": "2.1.3 (LlamaIndex)",
        "status": "running"
    }

@app.get("/health")
async def health_check():
    try:
        index_stats = pc_client.Index(config.PINECONE_INDEX_NAME).describe_index_stats()
        is_healthy = index_stats.status['ready']
        components = {
            'pinecone_index': is_healthy,
            'openai_api': True
        }
        if not is_healthy:
            logger.warning("⚠️ Pinecone index is not ready.")

        system_status = {
            'is_healthy': is_healthy,
            'components': components,
            'index_stats': index_stats.status,
            'document_count': index_stats.total_vector_count
        }
        
        logger.info(f"💚 Health check completed. Status: {'Healthy' if is_healthy else 'Unhealthy'}")
        return system_status
    except Exception as e:
        logger.error(f"❌ Health check failed: {e}")
        raise HTTPException(status_code=500, detail=f"Health check failed: {str(e)}")

# --- Document Processing Helpers ---
async def _extract_text_from_pdf(file_path: Path) -> List[str]:
    """Extracts text, tables, and images from a PDF file using PyMuPDF."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, _thread_process_pdf_file, file_path)

def _thread_process_pdf_file(file_path: Path) -> List[str]:
    """Synchronous PDF text, table, and image extraction for the thread pool."""
    all_pages_text = []
    try:
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            page_text = f"\n--- PAGE {page_num + 1} ---\n\n"
            
            # Extract tables and convert to markdown
            tables = page.find_tables()
            for table in tables:
                table_data = table.extract()
                # Use tabulate to create a markdown table representation
                table_md = tabulate(table_data, headers="firstrow", tablefmt="pipe")
                page_text += f"[TABLE_START]\n{table_md}\n[TABLE_END]\n\n"
            
            # Extract plain text
            page_text += page.get_text() or ""
            
            all_pages_text.append(page_text)
    except Exception as e:
        logger.error(f"Failed to extract content from PDF file {file_path}: {e}")
        return []
    return all_pages_text

async def _extract_text_from_docx(file_path: Path) -> List[str]:
    """Extracts text from a DOCX file."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, _thread_process_docx_file, file_path)

def _thread_process_docx_file(file_path: Path) -> List[str]:
    """Synchronous DOCX text extraction for the thread pool."""
    full_text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + '\n'
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX file {file_path}: {e}")
        return []
    return [full_text] # Return as a list for consistency

async def _extract_text_from_email(file_path: Path) -> List[str]:
    """Extracts text from an EML file."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, _thread_process_email_file, file_path)

def _thread_process_email_file(file_path: Path) -> List[str]:
    """Synchronous email text extraction for the thread pool."""
    full_text = ""
    try:
        with open(file_path, 'rb') as fp:
            msg = Message()
            msg.set_content(fp.read())
            if msg.is_multipart():
                for part in msg.walk():
                    ctype = part.get_content_type()
                    if ctype == 'text/plain':
                        full_text += part.get_payload(decode=True).decode() + '\n'
            else:
                full_text = msg.get_payload(decode=True).decode()
    except Exception as e:
        logger.error(f"Failed to extract text from email file {file_path}: {e}")
        return []
    return [full_text] # Return as a list for consistency

async def _load_and_index_from_url(document_url: str):
    """Helper function to load and index a document from a URL."""
    global index, pipeline
    if not pipeline:
        raise RuntimeError("Ingestion pipeline is not initialized.")
    logger.info(f"🌐 Downloading document from URL: {document_url}")
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        start_download = time.time()
        response = requests.get(document_url, timeout=60)
        response.raise_for_status()
        
        # Corrected file name parsing to handle invalid URL characters
        parsed_url = urlparse(document_url)
        clean_filename = os.path.basename(parsed_url.path)
        
        if not clean_filename:
            clean_filename = "document.pdf"
            
        temp_path = Path(temp_dir) / clean_filename
        temp_path.write_bytes(response.content)
        logger.info(f"📁 Document downloaded in {time.time() - start_download:.2f}s.")
        
        start_processing = time.time()
        
        file_extension = temp_path.suffix.lower()
        if file_extension == '.pdf':
            extractor = _extract_text_from_pdf
        elif file_extension == '.docx':
            extractor = _extract_text_from_docx
        elif file_extension == '.eml':
            extractor = _extract_text_from_email
        else:
            raise ValueError(f"Unsupported document type: {file_extension}")
        
        # This now returns a list of strings
        list_of_texts = await extractor(temp_path)
        
        if not list_of_texts:
            raise ValueError("Text extraction from document failed or returned empty content.")

        logger.info(f"✅ Text extraction completed in {time.time() - start_processing:.2f}s. Total pages extracted: {len(list_of_texts)}")
        
        # Create a Document object for each page with enhanced metadata
        documents = [
            Document(text=text, metadata={"source": document_url, "page_number": i + 1})
            for i, text in enumerate(list_of_texts)
        ]
        
        start_ingestion = time.time()
        logger.info("📦 Starting LlamaIndex ingestion pipeline...")
        await pipeline.arun(documents=documents)
        logger.info(f"🎉 Ingestion pipeline completed in {time.time() - start_ingestion:.2f}s.")
        
        pinecone_vector_store = PineconeVectorStore(
            pinecone_client=pc_client,
            index_name=config.PINECONE_INDEX_NAME,
            environment=config.PINECONE_ENVIRONMENT
        )
        storage_context = StorageContext.from_defaults(vector_store=pinecone_vector_store)
        index = VectorStoreIndex.from_vector_store(pinecone_vector_store, storage_context=storage_context)
        
        index_stats = pc_client.Index(config.PINECONE_INDEX_NAME).describe_index_stats()
        logger.info(f"🎉 Successfully loaded and indexed document. Total vectors: {index_stats.total_vector_count}")
        return True
    except Exception as e:
        logger.error(f"❌ Error loading and indexing document from URL: {e}")
        raise
    finally:
        if Path(temp_dir).exists():
            for item in Path(temp_dir).iterdir():
                os.remove(item)
            os.rmdir(temp_dir)
            logger.info(f"🗑️ Deleted temporary directory: {temp_dir}")

# --- API Endpoints ---
@app.post("/load-and-index", tags=["Document Management"])
async def load_and_index_document(request: dict, token: str = Depends(verify_token)):
    """Loads and indexes a document from a URL into the Pinecone index."""
    document_url = request.get("document_url")
    if not document_url:
        raise HTTPException(status_code=400, detail="Missing 'document_url' in request body.")

    start_time = time.time()
    try:
        await _load_and_index_from_url(document_url)
        
        return {
            "message": "Document loaded and indexed successfully.",
            "url": document_url,
            "processing_time": time.time() - start_time
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading document: {str(e)}")

@app.post("/query", tags=["Querying"])
async def process_query(request: QueryRequest, token: str = Depends(verify_token)) -> DecisionResponse:
    """Processes a single query against the indexed documents and returns a structured decision."""
    global index
    if index is None:
        raise HTTPException(status_code=404, detail="No document index available. Please load a document first.")
    logger.info(f"➡️ Received single query: '{request.query}'")
    start_time = time.time()
    
    query_engine = index.as_query_engine(
        similarity_top_k=config.TOP_K_RESULTS,
        response_mode="compact"
    )

    try:
        decision_dict = await query_processor.process_query_for_decision(request.query, query_engine)
        end_time = time.time()
        
        logger.info(f"✅ Query processed successfully in {end_time - start_time:.2f}s.")
        
        return DecisionResponse(**decision_dict)

    except Exception as e:
        logger.error(f"❌ Error processing query: {e}")
        raise HTTPException(status_code=500, detail=f"Query processing failed: {str(e)}")

@app.post("/hackrx/run", tags=["Hackathon"])
async def process_batch_queries(request: BatchQueryRequest, token: str = Depends(verify_token)) -> HackRxResponse:
    """
    Handles both document loading and batch querying
    using LlamaIndex's powerful query engine and parallel processing capabilities.
    """
    global index
    start_time = time.time()

    if request.documents:
        logger.info(f"📥 A new document URL was provided. Loading and indexing: {request.documents}")
        try:
            await _load_and_index_from_url(request.documents)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error loading document from URL: {str(e)}")
    else:
        try:
            index_stats = pc_client.Index(config.PINECONE_INDEX_NAME).describe_index_stats()
            if index_stats.total_vector_count == 0:
                raise HTTPException(status_code=404, detail="No document index available. Please use the 'documents' field to provide a URL for a document to be indexed.")
            logger.info(f"📂 Using existing document index with {index_stats.total_vector_count} vectors.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not connect to Pinecone index: {str(e)}")

    logger.info(f"➡️ Processing a batch of {len(request.questions)} questions.")

    if index is None:
        pinecone_vector_store = PineconeVectorStore(
            pinecone_client=pc_client,
            index_name=config.PINECONE_INDEX_NAME,
            environment=config.PINECONE_ENVIRONMENT
        )
        storage_context = StorageContext.from_defaults(vector_store=pinecone_vector_store)
        index = VectorStoreIndex.from_vector_store(pinecone_vector_store, storage_context=storage_context)

    query_engine = index.as_query_engine(
        similarity_top_k=config.TOP_K_RESULTS,
        response_mode="compact"
    )

    async def process_single_question(question: str):
        logger.info(f"   - Starting query for question: '{question}'")
        try:
            decision_response = await query_processor.process_query_for_decision(question, query_engine)
            logger.info(f"   - Finished query for question: '{question}'")
            return decision_response['justification']
        except Exception as e:
            logger.error(f"   - Failed to process question '{question}': {e}")
            return "Error processing this question."

    tasks = [process_single_question(q) for q in request.questions]
    answers = await asyncio.gather(*tasks)

    processing_time = time.time() - start_time
    logger.info(f"🎉 Batch query processing completed for {len(request.questions)} questions in {processing_time:.2f}s.")
    
    return HackRxResponse(answers=answers)
    
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=port,
        reload=False,
        log_level=config.LOG_LEVEL.lower()
    )